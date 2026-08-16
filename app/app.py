"""
Uzbek Text Simplifier — Tkinter desktop app.

Модель грузится напрямую с Hugging Face Hub (репозиторий MODEL_ID),
поэтому веса заливать вручную не нужно — при первом запуске
они скачаются автоматически и закешируются.

Запуск локально:
    pip install -r requirements.txt
    python app.py
"""

import os
import re
import time
import threading
import tkinter as tk
import tkinter.font as tkfont
from tkinter import ttk, messagebox, filedialog

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    _DND_AVAILABLE = True
    _BaseTk = TkinterDnD.Tk
except ImportError:
    _DND_AVAILABLE = False
    _BaseTk = tk.Tk

try:
    import docx as _docx_lib  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# Детерминированный, чисто-скриптованный (regex + словарь) конвертер
# упрощённого текста в Markdown. НЕ модель, НЕ API — обычная функция
# строка -> строка, выполняется мгновенно и синхронно.
from markdown_formatter import to_markdown
from config import SimplifierConfig, build_prefix

MODEL_ID = "ismailoviskandar02/uzbek-text-simplifier"  # твой репозиторий модели на HF
MODEL_SUBFOLDER = "model2"
# Локальная копия модели: если скачана (кнопка "Скачать модель"), веса
# грузятся с диска, а не с Hugging Face Hub при каждом запуске.
LOCAL_MODEL_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "model_local"
)
LOCAL_MODEL_SUBPATH = os.path.join(LOCAL_MODEL_DIR, MODEL_SUBFOLDER)
# PREFIX больше не хардкодится: используется build_prefix(current_config).
# Оставлен как fallback-константа на случай прямого импорта модуля где-то ещё.
PREFIX = "simplify: "
MAX_INPUT_LEN = 256
MAX_NEW_TOKENS = 256


def _local_model_available() -> bool:
    """Есть ли на диске уже скачанная копия модели (папка непустая)."""
    return os.path.isdir(LOCAL_MODEL_SUBPATH) and bool(os.listdir(LOCAL_MODEL_SUBPATH))

# ---------------------------------------------------------------------------
# Локализация (RU / UZ / EN)
# ---------------------------------------------------------------------------

LANGS = {
    "ru": {
        "flag": "🇷🇺",
        "name": "Русский",
        "title": "🇺🇿 Oʻzbek matnini soddalashtirish",
        "subtitle": "Упрощение узбекского текста с помощью нейросети",
        "description": (
            "Переводит сложный текст (юридический, официальный, государственный "
            "документ) в понятный, простой узбекский язык."
        ),
        "input_label": "Исходный текст",
        "input_placeholder": "Вставьте сложный узбекский текст сюда...",
        "submit": "✨ Упростить",
        "clear": "🗑️ Очистить",
        "open_file": "📂 Открыть файл",
        "download_docx": "⬇️ Скачать",
        "dnd_hint": "Перетащите .docx или .txt файл сюда, или",
        "output_label": "Упрощённый текст",
        "footer": "Сделано с ❤️ для узбекского языка",
        "loading": "Загрузка модели...",
        "ready": "Модель загружена.",
        "ready_local": "Модель загружена (локальная копия).",
        "simplifying": "Упрощение...",
        "download_model": "💾 Скачать модель",
        "downloading_model": "Скачивание модели на диск...",
        "model_downloaded": "Модель скачана. Перезагрузка из локальной копии...",
        "model_download_error": "Не удалось скачать модель",
        "model_already_local": "Модель уже скачана на диск.",
        "error_empty": "Введите текст для упрощения.",
        "error_model": "Модель ещё не загружена, подождите.",
        "error_file_read": "Не удалось прочитать файл",
        "error_file_type": "Поддерживаются только файлы .docx и .txt",
        "error_no_output": "Сначала получите упрощённый текст.",
        "error_docx_lib": "Для сохранения в .docx нужна библиотека python-docx (pip install python-docx).",
        "saved_ok": "Файл сохранён",
        "to_markdown": "📝 В Markdown",
        "to_markdown_undo": "↩️ Обычный текст",
        "logs_tip": "🧾 Логи",
        "theme_to_light": "☀️ Светлая",
        "theme_to_dark": "🌙 Тёмная",
        "log_window_title": "Логи обработки",
        "log_clear": "🗑️ Очистить логи",
        "log_copy": "📋 Копировать",
        "log_copied": "Логи скопированы в буфер обмена",
        "log_empty": "Пока нет ни одной записи.",
        "log_close": "Закрыть",
        "cfg_title": "Параметры упрощения",
        "cfg_aggressiveness": "Агрессивность",
        "cfg_agg_conservative": "Мягкая",
        "cfg_agg_balanced": "Сбалансированная",
        "cfg_agg_aggressive": "Сильная",
        "cfg_max_length": "Макс. длина результата",
        "cfg_beams": "Точность (num_beams)",
        "cfg_drop_dates": "Убирать даты",
        "cfg_drop_law_refs": "Убирать ссылки на законы",
        "cfg_drop_stats": "Убирать статистику/цифры",
    },
    "uz": {
        "flag": "🇺🇿",
        "name": "Oʻzbek tili",
        "title": "🇺🇿 Oʻzbek matnini soddalashtirish",
        "subtitle": "Sun'iy intellekt yordamida matnni soddalashtirish",
        "description": (
            "Murakkab matnni (huquqiy, rasmiy, davlat hujjati) tushunarli, "
            "oddiy oʻzbek tiliga oʻgiradi."
        ),
        "input_label": "Asl matn",
        "input_placeholder": "Murakkab oʻzbek matnini shu yerga joylashtiring...",
        "submit": "✨ Soddalashtirish",
        "clear": "🗑️ Tozalash",
        "open_file": "📂 Faylni ochish",
        "download_docx": "⬇️ Yuklab olish",
        "dnd_hint": "Bu yerga .docx yoki .txt faylni tashlang, yoki",
        "output_label": "Soddalashtirilgan matn",
        "footer": "Oʻzbek tili uchun ❤️ bilan yaratilgan",
        "loading": "Model yuklanmoqda...",
        "ready": "Model yuklandi.",
        "ready_local": "Model yuklandi (lokal nusxa).",
        "simplifying": "Soddalashtirilmoqda...",
        "download_model": "💾 Modelni yuklab olish",
        "downloading_model": "Model diskka yuklanmoqda...",
        "model_downloaded": "Model yuklandi. Lokal nusxadan qayta yuklanmoqda...",
        "model_download_error": "Modelni yuklab bo'lmadi",
        "model_already_local": "Model allaqachon diskda mavjud.",
        "error_empty": "Soddalashtirish uchun matn kiriting.",
        "error_model": "Model hali yuklanmagan, kuting.",
        "error_file_read": "Faylni o'qib bo'lmadi",
        "error_file_type": "Faqat .docx va .txt fayllar qo'llab-quvvatlanadi",
        "error_no_output": "Avval soddalashtirilgan matnni oling.",
        "error_docx_lib": ".docx saqlash uchun python-docx kutubxonasi kerak (pip install python-docx).",
        "saved_ok": "Fayl saqlandi",
        "to_markdown": "📝 Markdown'ga",
        "to_markdown_undo": "↩️ Oddiy matn",
        "logs_tip": "🧾 Loglar",
        "theme_to_light": "☀️ Yorugʻ",
        "theme_to_dark": "🌙 Tungi",
        "log_window_title": "Qayta ishlash loglari",
        "log_clear": "🗑️ Loglarni tozalash",
        "log_copy": "📋 Nusxalash",
        "log_copied": "Loglar xotiraga nusxalandi",
        "log_empty": "Hozircha yozuvlar yo'q.",
        "log_close": "Yopish",
        "cfg_title": "Soddalashtirish parametrlari",
        "cfg_aggressiveness": "Agressivlik",
        "cfg_agg_conservative": "Yumshoq",
        "cfg_agg_balanced": "Muvozanatli",
        "cfg_agg_aggressive": "Kuchli",
        "cfg_max_length": "Natija maks. uzunligi",
        "cfg_beams": "Aniqlik (num_beams)",
        "cfg_drop_dates": "Sanalarni olib tashlash",
        "cfg_drop_law_refs": "Qonun havolalarini olib tashlash",
        "cfg_drop_stats": "Statistika/raqamlarni olib tashlash",
    },
    "en": {
        "flag": "🇬🇧",
        "name": "English",
        "title": "🇺🇿 Uzbek Text Simplifier",
        "subtitle": "Simplify Uzbek text with AI",
        "description": (
            "Turns complex text (legal, official, government documents) "
            "into clear, simple Uzbek."
        ),
        "input_label": "Original text",
        "input_placeholder": "Paste complex Uzbek text here...",
        "submit": "✨ Simplify",
        "clear": "🗑️ Clear",
        "open_file": "📂 Open file",
        "download_docx": "⬇️ Download",
        "dnd_hint": "Drop a .docx or .txt file here, or",
        "output_label": "Simplified text",
        "footer": "Made with ❤️ for the Uzbek language",
        "loading": "Loading model...",
        "ready": "Model loaded.",
        "ready_local": "Model loaded (local copy).",
        "simplifying": "Simplifying...",
        "download_model": "💾 Download model",
        "downloading_model": "Downloading model to disk...",
        "model_downloaded": "Model downloaded. Reloading from local copy...",
        "model_download_error": "Failed to download the model",
        "model_already_local": "Model is already downloaded.",
        "error_empty": "Please enter text to simplify.",
        "error_model": "Model is still loading, please wait.",
        "error_file_read": "Could not read the file",
        "error_file_type": "Only .docx and .txt files are supported",
        "error_no_output": "Simplify some text first.",
        "error_docx_lib": "Saving .docx requires the python-docx package (pip install python-docx).",
        "saved_ok": "File saved",
        "to_markdown": "📝 To Markdown",
        "to_markdown_undo": "↩️ Plain text",
        "logs_tip": "🧾 Logs",
        "theme_to_light": "☀️ Light",
        "theme_to_dark": "🌙 Dark",
        "log_window_title": "Processing logs",
        "log_clear": "🗑️ Clear logs",
        "log_copy": "📋 Copy",
        "log_copied": "Logs copied to clipboard",
        "log_empty": "No entries yet.",
        "log_close": "Close",
        "cfg_title": "Simplification parameters",
        "cfg_aggressiveness": "Aggressiveness",
        "cfg_agg_conservative": "Conservative",
        "cfg_agg_balanced": "Balanced",
        "cfg_agg_aggressive": "Aggressive",
        "cfg_max_length": "Max output length",
        "cfg_beams": "Accuracy (num_beams)",
        "cfg_drop_dates": "Drop dates",
        "cfg_drop_law_refs": "Drop legal references",
        "cfg_drop_stats": "Drop statistics/numbers",
    },
}

# ---------------------------------------------------------------------------
# Темы (тёмная / светлая). Круглый дизайн опирается на эти же цвета —
# радиусы скругления заданы отдельно константами RADIUS_*.
# ---------------------------------------------------------------------------

THEMES = {
    "dark": {
        "bg": "#05070d",
        "panel": "#0b0f1a",
        "panel2": "#0f1524",
        "border": "#1e2740",
        "label_bg": "#141b2e",
        "label_fg": "#8fb3ff",
        "title_fg": "#e6edff",
        "text_fg": "#dbe4ff",
        "subdued_fg": "#7d8bb3",
        "input_bg": "#0d1220",
        "input_border": "#243257",
        "accent": "#3b6bff",
        "accent_hover": "#5b8bff",
        "accent_fg": "#ffffff",
        "btn_secondary_bg": "#141b2e",
        "btn_secondary_hover": "#1c2540",
        "btn_secondary_fg": "#c3cdf0",
        "danger": "#3a1f2b",
        "danger_hover": "#48283a",
        "danger_fg": "#ffb4c6",
    },
    "light": {
        "bg": "#f3f5fb",
        "panel": "#ffffff",
        "panel2": "#eef1fb",
        "border": "#dbe1f5",
        "label_bg": "#e9edfb",
        "label_fg": "#3552b8",
        "title_fg": "#1b2444",
        "text_fg": "#222a44",
        "subdued_fg": "#66709a",
        "input_bg": "#ffffff",
        "input_border": "#c9d3f0",
        "accent": "#3b6bff",
        "accent_hover": "#2a56e0",
        "accent_fg": "#ffffff",
        "btn_secondary_bg": "#e9edfb",
        "btn_secondary_hover": "#dce2f7",
        "btn_secondary_fg": "#2c3866",
        "danger": "#ffe4ea",
        "danger_hover": "#ffd2dd",
        "danger_fg": "#b3264a",
    },
}

RADIUS_PANEL = 18
RADIUS_BTN = 14
RADIUS_PILL = 999  # эффективно даёт полностью скруглённые "таблетки"


# ---------------------------------------------------------------------------
# Логи обработки — простая шина событий, которую читает LogWindow.
# ---------------------------------------------------------------------------

class ProcessingLog:
    """Хранит хронологию событий приложения (загрузка модели, упрощение,
    экспорт, ошибки и т.д.) и уведомляет подписчиков (открытое окно логов)
    о каждой новой записи."""

    def __init__(self):
        self.entries = []  # list[(time_str, level, message)]
        self._listeners = []

    def subscribe(self, callback):
        self._listeners.append(callback)

    def unsubscribe(self, callback):
        if callback in self._listeners:
            self._listeners.remove(callback)

    def add(self, message, level="INFO"):
        ts = time.strftime("%H:%M:%S")
        entry = (ts, level, message)
        self.entries.append(entry)
        for cb in list(self._listeners):
            try:
                cb(entry)
            except Exception:  # noqa: BLE001
                pass

    def clear(self):
        self.entries.clear()
        for cb in list(self._listeners):
            try:
                cb(None)
            except Exception:  # noqa: BLE001
                pass

    def as_text(self):
        return "\n".join(f"[{ts}] {lvl}: {msg}" for ts, lvl, msg in self.entries)


# ---------------------------------------------------------------------------
# Скруглённые виджеты. Tkinter не умеет border-radius нативно, поэтому
# кнопки и контейнеры полей рисуются на Canvas скруглёнными многоугольниками
# (create_polygon(..., smooth=True)), что даёт полноценный "круглый" дизайн
# без сторонних библиотек.
# ---------------------------------------------------------------------------

def _round_rect_points(x1, y1, x2, y2, r):
    return [
        x1 + r, y1, x2 - r, y1, x2, y1, x2, y1 + r,
        x2, y2 - r, x2, y2, x2 - r, y2, x1 + r, y2,
        x1, y2, x1, y2 - r, x1, y1 + r, x1, y1,
    ]


class RoundedButton(tk.Canvas):
    """Скруглённая кнопка на Canvas с поддержкой hover/disabled-состояний
    и произвольной палитрой (используется для переключения тёмной/светлой
    темы на лету через set_colors)."""

    def __init__(self, master, text="", command=None, bg_outer="#ffffff",
                 fill="#3b6bff", hover="#5b8bff", fg="#ffffff",
                 font=("Segoe UI", 11, "bold"), radius=RADIUS_BTN,
                 padx=18, pady=10, disabled_fill=None, disabled_fg=None,
                 min_width=0, **kw):
        super().__init__(master, highlightthickness=0, bd=0, bg=bg_outer, **kw)
        self.command = command
        self.fill = fill
        self.hover = hover
        self.fg = fg
        self.font = font
        self.radius = radius
        self.padx = padx
        self.pady = pady
        self.disabled_fill = disabled_fill or fill
        self.disabled_fg = disabled_fg or fg
        self.text_value = text
        self.state_ = "normal"
        self._hovering = False
        self.min_width = min_width

        self._recompute_size()
        self.configure(width=self._desired_w, height=self._desired_h)

        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self.bind("<Button-1>", self._on_click)
        self.bind("<Configure>", lambda e: self._redraw())
        self._redraw()

    def _recompute_size(self):
        f = tkfont.Font(font=self.font)
        lines = self.text_value.split("\n") if self.text_value else [""]
        tw = max((f.measure(ln) for ln in lines), default=0)
        th = f.metrics("linespace") * max(len(lines), 1)
        self._desired_w = max(tw + 2 * self.padx, self.min_width)
        self._desired_h = th + 2 * self.pady

    def _redraw(self):
        self.delete("all")
        w = self.winfo_width() or self._desired_w
        h = self.winfo_height() or self._desired_h
        r = max(4, min(self.radius, w // 2, h // 2))
        if self.state_ == "disabled":
            color, fg = self.disabled_fill, self.disabled_fg
        else:
            color = self.hover if self._hovering else self.fill
            fg = self.fg
        self.create_polygon(
            _round_rect_points(1, 1, max(w - 1, 3), max(h - 1, 3), r),
            smooth=True, splinesteps=24, fill=color, outline=color,
        )
        self.create_text(
            w / 2, h / 2, text=self.text_value, fill=fg, font=self.font,
            justify="center",
        )

    def _on_enter(self, _e=None):
        if self.state_ == "normal":
            self._hovering = True
            self.configure(cursor="hand2")
            self._redraw()

    def _on_leave(self, _e=None):
        self._hovering = False
        self.configure(cursor="")
        self._redraw()

    def _on_click(self, _e=None):
        if self.state_ == "normal" and self.command:
            self.command()

    def config_text(self, text):
        self.text_value = text
        self._recompute_size()
        self.configure(width=self._desired_w, height=self._desired_h)
        self._redraw()

    def config_state(self, state):
        self.state_ = state
        self._redraw()

    def set_colors(self, bg_outer=None, fill=None, hover=None, fg=None,
                    disabled_fill=None, disabled_fg=None):
        if bg_outer is not None:
            self.configure(bg=bg_outer)
        if fill is not None:
            self.fill = fill
        if hover is not None:
            self.hover = hover
        if fg is not None:
            self.fg = fg
        if disabled_fill is not None:
            self.disabled_fill = disabled_fill
        if disabled_fg is not None:
            self.disabled_fg = disabled_fg
        self._redraw()


class RoundedTextHost(tk.Canvas):
    """Скруглённый контейнер-подложка для tk.Text (поле ввода/вывода),
    визуально даёт эффект "круглого" текстового поля."""

    def __init__(self, master, child_factory, bg_outer, fill, border,
                 radius=RADIUS_PANEL, pad=12, **kw):
        super().__init__(master, highlightthickness=0, bd=0, bg=bg_outer, **kw)
        self.radius = radius
        self.fill = fill
        self.border = border
        self.pad = pad
        self.child = child_factory(self)
        self._win = self.create_window(pad, pad, anchor="nw", window=self.child)
        self.bind("<Configure>", lambda e: self.redraw())

    def redraw(self):
        w = self.winfo_width()
        h = self.winfo_height()
        if w < 4 or h < 4:
            return
        self.delete("bg")
        r = max(4, min(self.radius, w // 2, h // 2))
        self.create_polygon(
            _round_rect_points(1, 1, w - 1, h - 1, r),
            smooth=True, splinesteps=24, fill=self.fill,
            outline=self.border, width=1, tags=("bg",),
        )
        self.tag_lower("bg")
        cw = max(w - 2 * self.pad, 10)
        ch = max(h - 2 * self.pad, 10)
        self.coords(self._win, self.pad, self.pad)
        self.itemconfig(self._win, width=cw, height=ch)

    def set_colors(self, bg_outer=None, fill=None, border=None):
        if bg_outer is not None:
            self.configure(bg=bg_outer)
        if fill is not None:
            self.fill = fill
        if border is not None:
            self.border = border
        self.redraw()


def _make_readonly_text(widget):
    """Разрешает выделение и копирование текста (на ЛЮБОМ языке/скрипте —
    выделение в Tk работает на уровне символов и не зависит от алфавита),
    но блокирует редактирование. В отличие от state="disabled", это не
    ломает выделение мышью и стандартные комбинации клавиш."""
    allowed_keysyms = {
        "Left", "Right", "Up", "Down", "Home", "End", "Prior", "Next",
        "Shift_L", "Shift_R", "Control_L", "Control_R", "Tab",
    }

    def _on_key(event):
        ctrl = bool(event.state & 0x4)
        if ctrl and event.keysym.lower() in ("c", "a", "insert"):
            return None
        if event.keysym in allowed_keysyms:
            return None
        return "break"

    def _select_all(_event=None):
        widget.tag_add("sel", "1.0", "end-1c")
        widget.mark_set("insert", "end-1c")
        widget.see("insert")
        return "break"

    widget.bind("<Key>", _on_key)
    widget.bind("<Control-a>", _select_all)
    widget.bind("<Control-A>", _select_all)
    # запрещаем drag&drop-вставку и стандартную вставку из буфера
    widget.bind("<<Paste>>", lambda e: "break")
    widget.bind("<Button-2>", lambda e: "break")  # средняя кнопка (X11 paste)


# ---------------------------------------------------------------------------
# Окно логов обработки
# ---------------------------------------------------------------------------

class LogWindow(tk.Toplevel):
    def __init__(self, master, log: ProcessingLog, colors, lang_code):
        super().__init__(master)
        self.log = log
        self.colors = colors
        self.lang_code = lang_code
        self.withdraw()
        self.transient(master)
        self.configure(bg=colors["bg"])
        self.geometry("640x420")
        self.minsize(420, 260)

        self._build_ui()
        self._reload_all()
        self.log.subscribe(self._on_entry)
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self.deiconify()

    def _t(self):
        return LANGS[self.lang_code]

    def _build_ui(self):
        t = self._t()
        self.title(t["log_window_title"])

        root = tk.Frame(self, bg=self.colors["bg"], padx=14, pady=14)
        root.pack(fill="both", expand=True)

        header = tk.Label(
            root, text=t["log_window_title"], bg=self.colors["bg"],
            fg=self.colors["title_fg"], font=("Segoe UI", 14, "bold"),
        )
        header.pack(anchor="w", pady=(0, 8))

        host = RoundedTextHost(
            root,
            child_factory=lambda c: tk.Text(
                c, wrap="word", bg=self.colors["input_bg"],
                fg=self.colors["text_fg"], relief="flat",
                highlightthickness=0, font=("Consolas", 11),
                padx=4, pady=4,
            ),
            bg_outer=self.colors["bg"], fill=self.colors["input_bg"],
            border=self.colors["input_border"], radius=RADIUS_PANEL,
        )
        host.pack(fill="both", expand=True)
        self.text_widget = host.child
        self._host = host
        _make_readonly_text(self.text_widget)

        btn_row = tk.Frame(root, bg=self.colors["bg"])
        btn_row.pack(fill="x", pady=(10, 0))

        self.clear_btn = RoundedButton(
            btn_row, text=t["log_clear"], command=self._on_clear,
            bg_outer=self.colors["bg"], fill=self.colors["btn_secondary_bg"],
            hover=self.colors["btn_secondary_hover"], fg=self.colors["btn_secondary_fg"],
            font=("Segoe UI", 10, "bold"), radius=RADIUS_BTN,
        )
        self.clear_btn.pack(side="left")

        self.copy_btn = RoundedButton(
            btn_row, text=t["log_copy"], command=self._on_copy,
            bg_outer=self.colors["bg"], fill=self.colors["accent"],
            hover=self.colors["accent_hover"], fg=self.colors["accent_fg"],
            font=("Segoe UI", 10, "bold"), radius=RADIUS_BTN,
        )
        self.copy_btn.pack(side="left", padx=(8, 0))

        self.close_btn = RoundedButton(
            btn_row, text=t["log_close"], command=self._on_close,
            bg_outer=self.colors["bg"], fill=self.colors["btn_secondary_bg"],
            hover=self.colors["btn_secondary_hover"], fg=self.colors["btn_secondary_fg"],
            font=("Segoe UI", 10, "bold"), radius=RADIUS_BTN,
        )
        self.close_btn.pack(side="right")

    def _reload_all(self):
        t = self._t()
        self.text_widget.config(state="normal")
        self.text_widget.delete("1.0", "end")
        if not self.log.entries:
            self.text_widget.insert("1.0", t["log_empty"])
        else:
            self.text_widget.insert("1.0", self.log.as_text())
        self.text_widget.see("end")

    def _on_entry(self, entry):
        if entry is None:
            self.after(0, self._reload_all)
            return

        def append():
            ts, lvl, msg = entry
            self.text_widget.config(state="normal")
            content = self.text_widget.get("1.0", "end-1c")
            if content.strip() == self._t()["log_empty"]:
                self.text_widget.delete("1.0", "end")
                prefix = ""
            else:
                prefix = "\n" if content else ""
            self.text_widget.insert("end", f"{prefix}[{ts}] {lvl}: {msg}")
            self.text_widget.see("end")

        self.after(0, append)

    def _on_clear(self):
        self.log.clear()

    def _on_copy(self):
        self.clipboard_clear()
        self.clipboard_append(self.log.as_text())
        self.update()

    def apply_theme(self, colors):
        self.colors = colors
        self.configure(bg=colors["bg"])
        for w in self.winfo_children():
            w.configure(bg=colors["bg"]) if isinstance(w, tk.Frame) else None
        self._host.set_colors(bg_outer=colors["bg"], fill=colors["input_bg"], border=colors["input_border"])
        self.text_widget.configure(bg=colors["input_bg"], fg=colors["text_fg"])
        self.clear_btn.set_colors(
            bg_outer=colors["bg"], fill=colors["btn_secondary_bg"],
            hover=colors["btn_secondary_hover"], fg=colors["btn_secondary_fg"],
        )
        self.copy_btn.set_colors(
            bg_outer=colors["bg"], fill=colors["accent"],
            hover=colors["accent_hover"], fg=colors["accent_fg"],
        )
        self.close_btn.set_colors(
            bg_outer=colors["bg"], fill=colors["btn_secondary_bg"],
            hover=colors["btn_secondary_hover"], fg=colors["btn_secondary_fg"],
        )

    def apply_language(self, lang_code):
        self.lang_code = lang_code
        t = self._t()
        self.title(t["log_window_title"])
        self.clear_btn.config_text(t["log_clear"])
        self.copy_btn.config_text(t["log_copy"])
        self.close_btn.config_text(t["log_close"])
        self._reload_all()

    def _on_close(self):
        self.log.unsubscribe(self._on_entry)
        self.master.log_window = None
        self.destroy()


class ModelRunner:
    """Loads and runs the Uzbek text simplification model in the background."""

    def __init__(self, on_loaded=None, on_error=None):
        self.tokenizer = None
        self.model = None
        self.device = "cpu"
        self.loaded = False
        self.source = None  # "local" | "hub", заполняется после загрузки
        self.on_loaded = on_loaded
        self.on_error = on_error

    def load_async(self):
        threading.Thread(target=self._load, daemon=True).start()

    def _load(self):
        try:
            import torch
            from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

            self.device = "cuda" if torch.cuda.is_available() else "cpu"

            # Если модель уже скачана на диск (см. download_local) —
            # грузим её оттуда, без обращения к Hugging Face Hub.
            if _local_model_available():
                source, kwargs = LOCAL_MODEL_SUBPATH, {}
                self.source = "local"
            else:
                source, kwargs = MODEL_ID, {"subfolder": MODEL_SUBFOLDER}
                self.source = "hub"

            self.tokenizer = AutoTokenizer.from_pretrained(source, **kwargs)
            self.model = AutoModelForSeq2SeqLM.from_pretrained(
                source, **kwargs
            ).to(self.device)
            self.model.eval()
            self.loaded = True
            if self.on_loaded:
                self.on_loaded()
        except Exception as e:  # noqa: BLE001
            if self.on_error:
                self.on_error(str(e))

    def download_local(self):
        """Скачивает веса модели (только subfolder=model2) на диск в
        LOCAL_MODEL_DIR, чтобы дальнейшие запуски грузили модель с
        компьютера, а не с Hugging Face Hub. Синхронный вызов — гонять
        в отдельном потоке должен вызывающий код."""
        from huggingface_hub import snapshot_download

        os.makedirs(LOCAL_MODEL_DIR, exist_ok=True)
        snapshot_download(
            repo_id=MODEL_ID,
            allow_patterns=[f"{MODEL_SUBFOLDER}/*"],
            local_dir=LOCAL_MODEL_DIR,
        )

    def simplify(self, text: str, num_beams: int = 4, prefix: str = PREFIX) -> str:
        import torch

        text = (text or "").strip()
        if not text:
            return ""
        inputs = self.tokenizer(
            prefix + text,
            return_tensors="pt",
            truncation=True,
            max_length=MAX_INPUT_LEN,
        ).to(self.device)

        with torch.no_grad():
            out = self.model.generate(
                **inputs,
                max_new_tokens=MAX_NEW_TOKENS,
                num_beams=num_beams,
            )

        return self.tokenizer.decode(out[0], skip_special_tokens=True)


class SimplifierApp(_BaseTk):
    def __init__(self):
        super().__init__()
        self.lang_code = "ru"
        self.theme_name = "dark"
        self.colors = THEMES[self.theme_name]

        self.title(LANGS[self.lang_code]["title"])
        self.geometry("1040x840")
        self.configure(bg=self.colors["bg"])
        self.minsize(780, 660)

        self.log = ProcessingLog()
        self.log_window = None

        self.runner = ModelRunner(on_loaded=self._model_loaded, on_error=self._model_error)

        # Окна настроек больше нет — конфиг фиксирован и подобран под
        # максимальную точность упрощения (см. config.py).
        self.current_config = SimplifierConfig.load()

        # Состояние тоггла кнопки "В Markdown": показываем сейчас
        # Markdown-версию текста в поле вывода или обычную.
        self._output_markdown_active = False
        self._output_plain_text = ""
        # Сырая markdown-строка (с **, #, > и т.д.) для сохранения в .md —
        # само поле вывода этих символов не хранит, оно рендерит их стилями.
        self._output_markdown_raw = ""

        self._round_widgets = []  # (widget, kind) для смены темы налету

        self._build_style()
        self._build_ui()
        self._apply_language()

        self.log.add("Приложение запущено")
        self.status_var.set(LANGS[self.lang_code]["loading"])
        self.log.add(f"Загрузка модели {MODEL_ID}...")
        self.runner.load_async()

    # ------------------------------------------------------------------
    # Styling (ttk-часть темы; canvas-виджеты обновляются отдельно в
    # _apply_theme_to_round_widgets)
    # ------------------------------------------------------------------
    def _build_style(self):
        c = self.colors
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        style.configure("TFrame", background=c["bg"])
        style.configure("Panel.TFrame", background=c["panel"])
        style.configure(
            "TLabel", background=c["bg"], foreground=c["text_fg"],
            font=("Segoe UI", 11),
        )
        style.configure(
            "Title.TLabel", background=c["bg"], foreground=c["title_fg"],
            font=("Segoe UI", 21, "bold"),
        )
        style.configure(
            "Subtitle.TLabel", background=c["bg"], foreground=c["label_fg"],
            font=("Segoe UI", 13, "bold"),
        )
        style.configure(
            "Desc.TLabel", background=c["bg"], foreground=c["subdued_fg"],
            font=("Segoe UI", 11), wraplength=900, justify="center",
        )
        style.configure(
            "FieldLabel.TLabel", background=c["bg"], foreground=c["label_fg"],
            font=("Segoe UI", 11, "bold"),
        )
        style.configure(
            "Status.TLabel", background=c["bg"], foreground=c["subdued_fg"],
            font=("Segoe UI", 10, "italic"),
        )
        style.configure(
            "Footer.TLabel", background=c["bg"], foreground=c["subdued_fg"],
            font=("Segoe UI", 10),
        )
        style.configure(
            "TScale", background=c["bg"], troughcolor=c["input_bg"],
        )
        style.configure(
            "TCombobox",
            fieldbackground=c["input_bg"],
            background=c["input_bg"],
            foreground=c["text_fg"],
            arrowcolor=c["text_fg"],
        )
        style.configure(
            "TCheckbutton", background=c["bg"], foreground=c["text_fg"],
            font=("Segoe UI", 10),
        )
        style.map(
            "TCheckbutton",
            background=[("active", c["bg"])],
            foreground=[("active", c["text_fg"])],
        )
        self.option_add("*TCombobox*Listbox.background", c["input_bg"])
        self.option_add("*TCombobox*Listbox.foreground", c["text_fg"])
        self.option_add("*TCombobox*Listbox.selectBackground", c["accent"])

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------
    def _build_ui(self):
        c = self.colors
        root = ttk.Frame(self, style="TFrame", padding=18)
        root.pack(fill="both", expand=True)

        # ---- верхняя панель: лог-кнопка в левом верхнем углу, тема и
        # выбор языка — справа ----
        top_bar = ttk.Frame(root, style="TFrame")
        top_bar.pack(fill="x")

        self.log_btn = self._round_btn(
            top_bar, kind="secondary", text=LANGS[self.lang_code]["logs_tip"],
            command=self._open_log_window, font=("Segoe UI", 10, "bold"),
        )
        self.log_btn.pack(side="left")

        lang_row = ttk.Frame(top_bar, style="TFrame")
        lang_row.pack(side="right")

        self.lang_var = tk.StringVar(value=self.lang_code)
        lang_display = [f"{v['flag']} {v['name']}" for v in LANGS.values()]
        self._lang_keys = list(LANGS.keys())

        self.lang_combo = ttk.Combobox(
            lang_row, values=lang_display, state="readonly", width=16
        )
        self.lang_combo.current(self._lang_keys.index(self.lang_code))
        self.lang_combo.pack(side="right")
        self.lang_combo.bind("<<ComboboxSelected>>", self._on_language_change)

        self.theme_btn = self._round_btn(
            lang_row, kind="secondary", text=LANGS[self.lang_code]["theme_to_light"],
            command=self._toggle_theme, font=("Segoe UI", 10, "bold"),
        )
        self.theme_btn.pack(side="right", padx=(0, 10))

        # Header
        self.title_label = ttk.Label(root, text="", style="Title.TLabel", anchor="center")
        self.title_label.pack(pady=(14, 2), fill="x")

        self.subtitle_label = ttk.Label(root, text="", style="Subtitle.TLabel", anchor="center")
        self.subtitle_label.pack(fill="x")

        self.desc_label = ttk.Label(root, text="", style="Desc.TLabel", anchor="center")
        self.desc_label.pack(pady=(4, 14), fill="x")

        # Main content: two columns
        content = ttk.Frame(root, style="TFrame")
        content.pack(fill="both", expand=True)
        content.columnconfigure(0, weight=1)
        content.columnconfigure(1, weight=1)
        content.rowconfigure(0, weight=1)

        # Left column: input
        left = ttk.Frame(content, style="TFrame")
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        self.input_label = ttk.Label(left, text="", style="FieldLabel.TLabel")
        self.input_label.pack(anchor="w")

        self.input_host = RoundedTextHost(
            left,
            child_factory=lambda cv: tk.Text(
                cv, height=12, wrap="word",
                bg=c["input_bg"], fg=c["text_fg"],
                insertbackground=c["text_fg"],
                relief="flat", highlightthickness=0,
                font=("Segoe UI", 12), padx=2, pady=2,
            ),
            bg_outer=c["bg"], fill=c["input_bg"], border=c["input_border"],
            radius=RADIUS_PANEL,
        )
        self.input_host.pack(fill="both", expand=True, pady=(6, 6))
        self.input_text = self.input_host.child
        self.input_text.bind("<Control-a>", self._select_all_input)
        self.input_text.bind("<Control-A>", self._select_all_input)
        self._round_widgets.append((self.input_host, "text_host"))

        if _DND_AVAILABLE:
            self.input_text.drop_target_register(DND_FILES)
            self.input_text.dnd_bind("<<Drop>>", self._on_file_drop)

        file_row = ttk.Frame(left, style="TFrame")
        file_row.pack(fill="x", pady=(0, 10))

        self.dnd_hint_label = ttk.Label(file_row, text="", style="Status.TLabel")
        self.dnd_hint_label.pack(side="left")
        if not _DND_AVAILABLE:
            self.dnd_hint_label.pack_forget()

        self.open_file_btn = self._round_btn(
            file_row, kind="secondary", text="", command=self._on_open_file,
        )
        self.open_file_btn.pack(side="left", padx=(6, 0))
        # Замеряем реальную высоту этой строки после раскладки — нужна
        # ниже для спейсера справа (см. вывод колонку), чтобы оба
        # текстовых поля были одинаковой высоты без "магических" чисел.
        file_row.update_idletasks()
        self._file_row_height = file_row.winfo_reqheight()

        # Right column: output
        right = ttk.Frame(content, style="TFrame")
        right.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        self.output_label = ttk.Label(right, text="", style="FieldLabel.TLabel")
        self.output_label.pack(anchor="w")

        self.output_host = RoundedTextHost(
            right,
            child_factory=lambda cv: tk.Text(
                cv, height=12, wrap="word",
                bg=c["input_bg"], fg=c["text_fg"],
                insertbackground=c["text_fg"],
                relief="flat", highlightthickness=0,
                font=("Segoe UI", 12), padx=2, pady=2,
            ),
            bg_outer=c["bg"], fill=c["input_bg"], border=c["input_border"],
            radius=RADIUS_PANEL,
        )
        self.output_host.pack(fill="both", expand=True, pady=(6, 10))
        self.output_text = self.output_host.child

        # Спейсер той же высоты, что file_row слева (подсказка drag&drop +
        # "Открыть файл"), — чтобы оба текстовых поля (ввод и вывод)
        # занимали одинаковую высоту и были выровнены по нижнему краю.
        # Обычный пустой Frame с реально замеренной высотой file_row —
        # никакой отрисовки поверх него, поэтому не может "просвечивать"
        # чужим цветом (в отличие от предыдущей версии на Canvas-кнопке).
        self.output_spacer = tk.Frame(right, bg=c["bg"], height=self._file_row_height)
        self.output_spacer.pack(fill="x", pady=(0, 10))
        self.output_spacer.pack_propagate(False)
        # Ключевое исправление читаемости: поле вывода остаётся "normal",
        # редактирование блокируется через _make_readonly_text — поэтому
        # выделение текста мышью работает КОРРЕКТНО на любом языке/скрипте
        # (кириллица, латиница, смешанный текст), в отличие от
        # state="disabled", которое в Tk блокирует визуальное выделение.
        _make_readonly_text(self.output_text)
        self._configure_markdown_tags()
        self._round_widgets.append((self.output_host, "text_host"))

        # ---- Панель параметров упрощения — между текстовыми полями и
        # кнопками действий, во всю ширину. ----
        self._build_config_panel(root)

        # ---- Кнопки действий: одна общая строка на всю ширину под
        # панелью параметров, слева — работа с исходным текстом, справа —
        # работа с результатом. ----
        actions_row = ttk.Frame(root, style="TFrame")
        actions_row.pack(fill="x", pady=(0, 10))
        actions_row.columnconfigure(0, weight=1)
        actions_row.columnconfigure(1, weight=1)

        btn_row = ttk.Frame(actions_row, style="TFrame")
        btn_row.grid(row=0, column=0, sticky="w", padx=(0, 10))

        self.clear_btn = self._round_btn(
            btn_row, kind="secondary", text="", command=self._on_clear,
        )
        self.clear_btn.pack(side="left", padx=(0, 8))

        self.submit_btn = self._round_btn(
            btn_row, kind="accent", text="", command=self._on_submit,
        )
        self.submit_btn.pack(side="left", padx=(0, 8))

        self.download_model_btn = self._round_btn(
            btn_row, kind="secondary", text="", command=self._on_download_model,
        )
        self.download_model_btn.pack(side="left")

        output_btn_row = ttk.Frame(actions_row, style="TFrame")
        output_btn_row.grid(row=0, column=1, sticky="e", padx=(10, 0))

        self.markdown_btn = self._round_btn(
            output_btn_row, kind="secondary", text="", command=self._on_to_markdown,
        )
        self.markdown_btn.pack(side="left", padx=(0, 8))

        self.download_btn = self._round_btn(
            output_btn_row, kind="secondary", text="", command=self._on_download_docx,
        )
        self.download_btn.pack(side="left")

        # Status + footer
        self.status_var = tk.StringVar(value="")
        self.status_label = ttk.Label(root, textvariable=self.status_var, style="Status.TLabel")
        self.status_label.pack(pady=(10, 0))

        self.footer_label = ttk.Label(root, text="", style="Footer.TLabel", anchor="center")
        self.footer_label.pack(pady=(8, 0), fill="x")

    def _build_config_panel(self, root):
        """Компактная панель параметров упрощения на главном экране
        (замена отдельного окна настроек). Меняет self.current_config
        "на лету" и сохраняет его на диск при каждом изменении —
        следующий вызов _on_submit сразу подхватывает новые значения."""
        c = self.colors
        t = LANGS[self.lang_code]

        panel = ttk.Frame(root, style="TFrame")
        panel.pack(fill="x", pady=(0, 16))
        self._config_panel = panel

        self.cfg_title_label = ttk.Label(panel, text="", style="FieldLabel.TLabel")
        self.cfg_title_label.pack(anchor="w", pady=(0, 8))

        row = ttk.Frame(panel, style="TFrame")
        row.pack(fill="x")

        # --- Агрессивность ---
        agg_box = ttk.Frame(row, style="TFrame")
        agg_box.pack(side="left", padx=(0, 24))
        self.cfg_agg_label = ttk.Label(agg_box, text="", style="Status.TLabel")
        self.cfg_agg_label.pack(anchor="w")
        self._agg_values = ["conservative", "balanced", "aggressive"]
        self.agg_combo = ttk.Combobox(
            agg_box, state="readonly", width=16,
            values=[t["cfg_agg_conservative"], t["cfg_agg_balanced"], t["cfg_agg_aggressive"]],
        )
        self.agg_combo.current(self._agg_values.index(self.current_config.aggressiveness))
        self.agg_combo.pack(anchor="w", pady=(2, 0))
        self.agg_combo.bind("<<ComboboxSelected>>", self._on_cfg_aggressiveness_change)

        # --- Макс. длина результата ---
        len_box = ttk.Frame(row, style="TFrame")
        len_box.pack(side="left", padx=(0, 24))
        self.cfg_len_label = ttk.Label(len_box, text="", style="Status.TLabel")
        self.cfg_len_label.pack(anchor="w")
        self.max_len_var = tk.DoubleVar(value=self.current_config.max_length_ratio)
        self.max_len_scale = ttk.Scale(
            len_box, from_=0.3, to=1.0, orient="horizontal", length=140,
            variable=self.max_len_var, command=self._on_cfg_max_length_change,
        )
        self.max_len_scale.pack(anchor="w", pady=(2, 0))

        # --- Точность (num_beams) ---
        beams_box = ttk.Frame(row, style="TFrame")
        beams_box.pack(side="left", padx=(0, 24))
        self.cfg_beams_label = ttk.Label(beams_box, text="", style="Status.TLabel")
        self.cfg_beams_label.pack(anchor="w")
        self.beams_var = tk.IntVar(value=self.current_config.num_beams)
        self.beams_scale = ttk.Scale(
            beams_box, from_=1, to=8, orient="horizontal", length=140,
            variable=self.beams_var, command=self._on_cfg_beams_change,
        )
        self.beams_scale.pack(anchor="w", pady=(2, 0))

        # --- Чекбоксы drop_* ---
        flags_box = ttk.Frame(row, style="TFrame")
        flags_box.pack(side="left")

        self.drop_dates_var = tk.BooleanVar(value=self.current_config.drop_dates)
        self.drop_dates_chk = ttk.Checkbutton(
            flags_box, text="", variable=self.drop_dates_var,
            command=self._on_cfg_flags_change,
        )
        self.drop_dates_chk.pack(anchor="w")

        self.drop_law_var = tk.BooleanVar(value=self.current_config.drop_law_refs)
        self.drop_law_chk = ttk.Checkbutton(
            flags_box, text="", variable=self.drop_law_var,
            command=self._on_cfg_flags_change,
        )
        self.drop_law_chk.pack(anchor="w")

        self.drop_stats_var = tk.BooleanVar(value=self.current_config.drop_stats)
        self.drop_stats_chk = ttk.Checkbutton(
            flags_box, text="", variable=self.drop_stats_var,
            command=self._on_cfg_flags_change,
        )
        self.drop_stats_chk.pack(anchor="w")

        self._refresh_config_panel_text()

    def _refresh_config_panel_text(self):
        """Обновляет подписи панели параметров при смене языка/темы."""
        t = LANGS[self.lang_code]
        self.cfg_title_label.config(text=t["cfg_title"])
        self.cfg_agg_label.config(text=t["cfg_aggressiveness"])
        self.cfg_len_label.config(
            text=f"{t['cfg_max_length']}: {self.current_config.max_length_ratio:.2f}"
        )
        self.cfg_beams_label.config(
            text=f"{t['cfg_beams']}: {self.current_config.num_beams}"
        )
        self.drop_dates_chk.config(text=t["cfg_drop_dates"])
        self.drop_law_chk.config(text=t["cfg_drop_law_refs"])
        self.drop_stats_chk.config(text=t["cfg_drop_stats"])
        agg_values = [t["cfg_agg_conservative"], t["cfg_agg_balanced"], t["cfg_agg_aggressive"]]
        self.agg_combo.config(values=agg_values)
        self.agg_combo.current(self._agg_values.index(self.current_config.aggressiveness))

    def _save_current_config(self):
        self.current_config.save()

    def _on_cfg_aggressiveness_change(self, _event=None):
        idx = self.agg_combo.current()
        self.current_config.aggressiveness = self._agg_values[idx]
        self._save_current_config()
        self.log.add(f"Настройки: aggressiveness={self.current_config.aggressiveness}")

    def _on_cfg_max_length_change(self, _value=None):
        self.current_config.max_length_ratio = round(self.max_len_var.get(), 2)
        t = LANGS[self.lang_code]
        self.cfg_len_label.config(
            text=f"{t['cfg_max_length']}: {self.current_config.max_length_ratio:.2f}"
        )
        self._save_current_config()

    def _on_cfg_beams_change(self, _value=None):
        beams = int(round(self.beams_var.get()))
        self.current_config.num_beams = beams
        t = LANGS[self.lang_code]
        self.cfg_beams_label.config(text=f"{t['cfg_beams']}: {beams}")
        self._save_current_config()

    def _on_cfg_flags_change(self):
        self.current_config.drop_dates = self.drop_dates_var.get()
        self.current_config.drop_law_refs = self.drop_law_var.get()
        self.current_config.drop_stats = self.drop_stats_var.get()
        self._save_current_config()
        self.log.add("Настройки: drop-флаги обновлены")

    # ------------------------------------------------------------------
    # Помощник создания круглых кнопок с учётом текущей темы + учёт для
    # последующего переключения темы налету.
    # ------------------------------------------------------------------
    def _round_btn(self, master, kind, text, command, font=("Segoe UI", 11, "bold"),
                    padx=18, min_width=0):
        c = self.colors
        if kind == "accent":
            fill, hover, fg = c["accent"], c["accent_hover"], c["accent_fg"]
        else:
            fill, hover, fg = c["btn_secondary_bg"], c["btn_secondary_hover"], c["btn_secondary_fg"]
        btn = RoundedButton(
            master, text=text, command=command, bg_outer=c["bg"],
            fill=fill, hover=hover, fg=fg, font=font, radius=RADIUS_BTN,
            padx=padx, min_width=min_width,
        )
        self._round_widgets.append((btn, kind))
        return btn

    def _select_all_input(self, _event=None):
        self.input_text.tag_add("sel", "1.0", "end-1c")
        return "break"

    # ------------------------------------------------------------------
    # Тема
    # ------------------------------------------------------------------
    def _toggle_theme(self):
        self.theme_name = "light" if self.theme_name == "dark" else "dark"
        self.colors = THEMES[self.theme_name]
        self.log.add(f"Тема переключена на: {self.theme_name}")
        self._apply_theme()

    def _apply_theme(self):
        c = self.colors
        self.configure(bg=c["bg"])
        self._build_style()

        for widget, kind in self._round_widgets:
            if kind == "text_host":
                widget.set_colors(bg_outer=c["bg"], fill=c["input_bg"], border=c["input_border"])
                widget.child.configure(bg=c["input_bg"], fg=c["text_fg"], insertbackground=c["text_fg"])
            elif kind == "accent":
                widget.set_colors(
                    bg_outer=c["bg"], fill=c["accent"], hover=c["accent_hover"], fg=c["accent_fg"],
                )
            elif kind == "example":
                widget.set_colors(bg_outer=c["bg"], fill=c["panel2"], hover=c["panel"], fg=c["subdued_fg"])
            else:  # secondary
                widget.set_colors(
                    bg_outer=c["bg"], fill=c["btn_secondary_bg"],
                    hover=c["btn_secondary_hover"], fg=c["btn_secondary_fg"],
                )

        self.output_spacer.configure(bg=c["bg"])

        self._configure_markdown_tags()
        # если в поле вывода сейчас отрисован markdown — перерисуем со
        # свежими тегами/цветами
        if self._output_markdown_active:
            self._render_markdown_into_output(self._output_markdown_raw)
        else:
            self._set_output(self._output_plain_text)

        self.theme_btn.config_text(
            LANGS[self.lang_code]["theme_to_light" if self.theme_name == "dark" else "theme_to_dark"]
        )

        if self.log_window is not None:
            self.log_window.apply_theme(c)

    # ------------------------------------------------------------------
    # Окно логов
    # ------------------------------------------------------------------
    def _open_log_window(self):
        if self.log_window is not None:
            self.log_window.lift()
            self.log_window.focus_force()
            return
        self.log_window = LogWindow(self, self.log, self.colors, self.lang_code)

    # ------------------------------------------------------------------
    # Language handling
    # ------------------------------------------------------------------
    def _on_language_change(self, event=None):
        idx = self.lang_combo.current()
        self.lang_code = self._lang_keys[idx]
        self.log.add(f"Язык интерфейса изменён на: {self.lang_code}")
        self._apply_language()
        if self.log_window is not None:
            self.log_window.apply_language(self.lang_code)

    def _apply_language(self):
        t = LANGS[self.lang_code]
        self.title(t["title"])
        self.title_label.config(text=t["title"])
        self.subtitle_label.config(text=t["subtitle"])
        self.desc_label.config(text=t["description"])
        self.input_label.config(text=t["input_label"])
        self.submit_btn.config_text(t["submit"])
        self.clear_btn.config_text(t["clear"])
        self.open_file_btn.config_text(t["open_file"])
        self.dnd_hint_label.config(text=t["dnd_hint"])
        self.download_btn.config_text(t["download_docx"])
        self.download_model_btn.config_text(t["download_model"])
        self.markdown_btn.config_text(
            t["to_markdown_undo"] if self._output_markdown_active else t["to_markdown"]
        )
        self.output_label.config(text=t["output_label"])
        self.footer_label.config(text=t["footer"])
        self.log_btn.config_text(t["logs_tip"])
        self.theme_btn.config_text(
            t["theme_to_light"] if self.theme_name == "dark" else t["theme_to_dark"]
        )
        self._refresh_config_panel_text()

        if self.runner.loaded:
            self.status_var.set(t["ready_local"] if self.runner.source == "local" else t["ready"])
        elif self.status_var.get():
            self.status_var.set(t["loading"])

    # ------------------------------------------------------------------
    # Model callbacks
    # ------------------------------------------------------------------
    def _model_loaded(self):
        def show():
            t = LANGS[self.lang_code]
            self.status_var.set(t["ready_local"] if self.runner.source == "local" else t["ready"])
            self.log.add(
                f"Модель загружена (устройство: {self.runner.device}, "
                f"источник: {self.runner.source})"
            )
        self.after(0, show)

    def _model_error(self, err):
        def show():
            self.status_var.set(f"Error: {err}")
            self.log.add(f"Ошибка загрузки модели: {err}", level="ERROR")
            messagebox.showerror("Model load error", err)
        self.after(0, show)

    # ------------------------------------------------------------------
    # Скачивание модели на диск ("💾 Скачать модель"): после успешной
    # загрузки модель перечитывается уже из локальной копии (LOCAL_MODEL_DIR),
    # см. ModelRunner._load / _local_model_available.
    # ------------------------------------------------------------------
    def _on_download_model(self):
        t = LANGS[self.lang_code]

        if _local_model_available():
            self.status_var.set(t["model_already_local"])
            self.log.add("Модель уже скачана локально — повторное скачивание не требуется")
            return

        self.download_model_btn.config_state("disabled")
        self.status_var.set(t["downloading_model"])
        self.log.add(f"Скачивание модели {MODEL_ID} на диск в {LOCAL_MODEL_DIR}...")

        def worker():
            try:
                self.runner.download_local()
            except Exception as e:  # noqa: BLE001
                err = str(e)
            else:
                err = None

            def finish():
                self.download_model_btn.config_state("normal")
                if err:
                    self.status_var.set(f"Error: {err}")
                    self.log.add(f"Ошибка скачивания модели: {err}", level="ERROR")
                    messagebox.showerror(t["title"], f"{t['model_download_error']}: {err}")
                    return
                self.status_var.set(t["model_downloaded"])
                self.log.add("Модель скачана локально, перезагрузка из локальной копии...")
                # Перечитываем модель — теперь _local_model_available() вернёт
                # True, и ModelRunner._load возьмёт веса с диска.
                self.runner.loaded = False
                self.runner.load_async()

            self.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()

    def _on_clear(self):
        self.input_text.delete("1.0", "end")
        self._set_output("")
        self._reset_markdown_toggle()
        self.log.add("Поля очищены")

    # ------------------------------------------------------------------
    # File import (drag & drop + open dialog)
    # ------------------------------------------------------------------
    def _read_text_file(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        if ext == ".docx":
            doc = _docx_lib.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError("unsupported_type")

    def _load_file_into_input(self, path: str):
        t = LANGS[self.lang_code]
        ext = os.path.splitext(path)[1].lower()
        if ext not in (".txt", ".docx"):
            self.log.add(f"Неподдерживаемый тип файла: {path}", level="ERROR")
            messagebox.showwarning(t["title"], t["error_file_type"])
            return
        if ext == ".docx" and not _DOCX_AVAILABLE:
            self.log.add("python-docx не установлен — .docx недоступен", level="ERROR")
            messagebox.showerror(t["title"], t["error_docx_lib"])
            return
        try:
            text = self._read_text_file(path)
        except Exception as e:  # noqa: BLE001
            self.log.add(f"Не удалось прочитать файл {path}: {e}", level="ERROR")
            messagebox.showerror(t["title"], f"{t['error_file_read']}: {e}")
            return
        self.input_text.delete("1.0", "end")
        self.input_text.insert("1.0", text)
        self.log.add(f"Файл загружен в поле ввода: {os.path.basename(path)}")

    def _on_open_file(self):
        path = filedialog.askopenfilename(
            filetypes=[("Text / Word", "*.txt *.docx"), ("All files", "*.*")]
        )
        if path:
            self._load_file_into_input(path)

    def _on_file_drop(self, event):
        # tkinterdnd2 gives a Tcl list; braces wrap paths containing spaces
        paths = self.tk.splitlist(event.data)
        if paths:
            self.log.add(f"Файл перетащен в окно: {paths[0]}")
            self._load_file_into_input(paths[0])

    # ------------------------------------------------------------------
    # File export (.docx или .md)
    # ------------------------------------------------------------------
    def _on_download_docx(self):
        t = LANGS[self.lang_code]

        # Источник истины — сохранённые строки, а не содержимое виджета:
        # в режиме Markdown сам виджет уже не хранит сырые **/`/# символы
        # (они сняты и заменены визуальными стилями), поэтому для .md
        # экспорта берём _output_markdown_raw, а для .docx/.txt — всегда
        # обычный текст без markdown-разметки.
        if self._output_markdown_active:
            text_for_md = self._output_markdown_raw.strip()
        else:
            text_for_md = self._output_plain_text.strip()
        text_for_docx = self._output_plain_text.strip()

        if not text_for_md:
            messagebox.showwarning(t["title"], t["error_no_output"])
            return

        # Если сейчас активен режим просмотра Markdown — по умолчанию
        # предлагаем сохранить именно в .md, иначе — в .docx. Пользователь
        # в любом случае может выбрать другой тип файла в самом диалоге.
        default_ext = ".md" if self._output_markdown_active else ".docx"
        default_name = "simplified.md" if self._output_markdown_active else "simplified.docx"

        path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=[
                ("Markdown", "*.md"),
                ("Word document", "*.docx"),
            ],
            initialfile=default_name,
        )
        if not path:
            return

        ext = os.path.splitext(path)[1].lower()

        try:
            if ext == ".md":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text_for_md)
            else:
                if not _DOCX_AVAILABLE:
                    self.log.add("python-docx не установлен — сохранение .docx недоступно", level="ERROR")
                    messagebox.showerror(t["title"], t["error_docx_lib"])
                    return
                doc = _docx_lib.Document()
                for line in text_for_docx.split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
        except Exception as e:  # noqa: BLE001
            self.log.add(f"Ошибка сохранения файла {path}: {e}", level="ERROR")
            messagebox.showerror(t["title"], f"{t['error_file_read']}: {e}")
            return
        self.status_var.set(f"{t['saved_ok']}: {os.path.basename(path)}")
        self.log.add(f"Файл сохранён: {os.path.basename(path)}")

    # ------------------------------------------------------------------
    # "В Markdown": детерминированное regex/словарное преобразование,
    # без модели и без API — выполняется мгновенно, синхронно.
    # Кнопка работает как переключатель (тоггл, "как радиокнопка"):
    # первый клик заменяет текст в поле вывода на Markdown-версию,
    # повторный клик возвращает обычный (исходный упрощённый) текст —
    # новое окно нигде не открывается.
    # ------------------------------------------------------------------
    def _on_to_markdown(self):
        t = LANGS[self.lang_code]

        if self._output_markdown_active:
            # Уже показываем Markdown -> вернуть обычный текст
            self._set_output(self._output_plain_text)
            self._output_markdown_active = False
            self.markdown_btn.config_text(t["to_markdown"])
            self.log.add("Переключено обратно на обычный текст")
            return

        text = self._output_plain_text.strip()
        if not text:
            messagebox.showwarning(t["title"], t["error_no_output"])
            return

        markdown_text = to_markdown(text)  # чистая функция строка -> строка
        self._output_markdown_active = True
        self._set_output(markdown_text, is_markdown=True)
        self.markdown_btn.config_text(t["to_markdown_undo"])
        self.log.add("Текст преобразован в Markdown")

    def _reset_markdown_toggle(self):
        """Сбрасывает состояние тоггла Markdown (вызывается при новом
        упрощении текста, очистке или загрузке нового файла)."""
        self._output_markdown_active = False
        self._output_markdown_raw = ""
        t = LANGS[self.lang_code]
        self.markdown_btn.config_text(t["to_markdown"])

    def _configure_markdown_tags(self):
        """Настраивает визуальные стили для рендера Markdown прямо в
        поле вывода: заголовки, **bold**, `code`, > blockquote, таблицы
        и горизонтальные разделители — без показа сырых спецсимволов."""
        c = self.colors
        w = self.output_text
        w.tag_configure(
            "h1", font=("Segoe UI", 19, "bold"),
            foreground=c["title_fg"], spacing1=10, spacing3=14,
        )
        w.tag_configure(
            "h2", font=("Segoe UI", 16, "bold"),
            foreground=c["title_fg"], spacing1=16, spacing3=12,
        )
        w.tag_configure(
            "h3", font=("Segoe UI", 13, "bold"),
            foreground=c["label_fg"], spacing1=12, spacing3=10,
        )
        w.tag_configure("bold", font=("Segoe UI", 12, "bold"))
        w.tag_configure(
            "code", font=("Consolas", 11),
            background=c["panel2"], foreground=c["label_fg"],
        )
        w.tag_configure(
            "blockquote", font=("Segoe UI", 12, "italic"),
            foreground=c["subdued_fg"], lmargin1=16, lmargin2=16,
            spacing1=6, spacing3=10,
        )
        w.tag_configure(
            "table", font=("Consolas", 11), foreground=c["text_fg"],
            spacing3=2,
        )
        w.tag_configure(
            "hr", foreground=c["input_border"], font=("Segoe UI", 9),
            spacing1=8, spacing3=8,
        )
        # обычный текст (абзацы)
        w.tag_configure("para", spacing1=2, spacing2=4, spacing3=14)
        # пункты списков — заметный отступ снизу между пунктами
        w.tag_configure("listitem", spacing1=2, spacing2=3, spacing3=8)
        w.tag_configure("sublistitem", spacing1=1, spacing2=2, spacing3=6)

    _INLINE_MD_RE = re.compile(r"(\*\*[^*\n]+\*\*|`[^`\n]+`)")

    def _insert_inline_markdown(self, text, extra_tags=()):
        """Вставляет строку в output_text, снимая **bold**/`code`
        спецсимволы и применяя вместо них визуальные теги."""
        w = self.output_text
        for part in self._INLINE_MD_RE.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) >= 4:
                w.insert("end", part[2:-2], (*extra_tags, "bold"))
            elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
                w.insert("end", part[1:-1], (*extra_tags, "code"))
            else:
                w.insert("end", part, extra_tags)

    def _render_markdown_into_output(self, md_text):
        """Рисует Markdown в самом поле вывода (не как обычный текст
        с ** и #, а со стилями): заголовки крупнее/жирным, **bold**,
        `code`, > цитаты с полосой слева, таблицы моноширинным
        шрифтом, --- разделителем.

        Для читаемости после КАЖДОГО содержательного элемента (заголовок,
        цитата, строка таблицы, пункт списка, обычная строка) вставляется
        настоящая пустая строка — а не только там, где в исходном
        markdown была пустая строка. Пустые строки самого markdown при
        этом не дублируются (не копим по 2-3 подряд)."""
        w = self.output_text
        w.delete("1.0", "end")

        prev_blank = True  # чтобы не начинать текст с пустой строки

        def _blank_line():
            nonlocal prev_blank
            if not prev_blank:
                w.insert("end", "\n")
                prev_blank = True

        for raw_line in md_text.split("\n"):
            line = raw_line.rstrip()
            s = line.strip()

            if s == "":
                # исходная пустая строка markdown -> не добавляем лишнего,
                # просто гарантируем, что разделение уже есть
                _blank_line()
                continue

            if s == "---":
                _blank_line()
                w.insert("end", "─" * 48 + "\n", ("hr",))
                _blank_line()
                continue

            if s.startswith("### "):
                _blank_line()
                self._insert_inline_markdown(s[4:], ("h3",))
                w.insert("end", "\n")
                _blank_line()
                continue

            if s.startswith("## "):
                _blank_line()
                self._insert_inline_markdown(s[3:], ("h2",))
                w.insert("end", "\n")
                _blank_line()
                continue

            if s.startswith("# "):
                _blank_line()
                self._insert_inline_markdown(s[2:], ("h1",))
                w.insert("end", "\n")
                _blank_line()
                continue

            if s.startswith("> "):
                w.insert("end", "▍ ", ("blockquote",))
                self._insert_inline_markdown(s[2:], ("blockquote",))
                w.insert("end", "\n")
                _blank_line()
                continue

            if s.startswith("| "):
                w.insert("end", line + "\n", ("table",))
                prev_blank = False
                continue

            # обычная строка / пункт списка / вложенный подпункт —
            # сохраняем исходный отступ (для a)/b) вложенности)
            leading = len(line) - len(line.lstrip(" "))
            is_list_item = bool(re.match(r"^(\d+\.|-)\s", s))
            is_sub_item = bool(re.match(r"^[a-zа-яё]\)\s", s, flags=re.IGNORECASE))
            tag = "sublistitem" if (leading and is_sub_item) else (
                "listitem" if is_list_item else "para"
            )
            if leading:
                w.insert("end", line[:leading])
            self._insert_inline_markdown(line[leading:], (tag,))
            w.insert("end", "\n")
            # между пунктами списка тоже пустая строка — читаемее длинных
            # плотных перечней
            _blank_line()

    def _set_output(self, text, is_markdown=False):
        if is_markdown:
            self._output_markdown_raw = text
            self._render_markdown_into_output(text)
            return
        self._output_plain_text = text
        w = self.output_text
        w.delete("1.0", "end")
        # Текст не меняем (никаких лишних \n) — визуальный "воздух" между
        # строками/абзацами даёт тег "para" (spacing3), а не сами символы.
        w.insert("1.0", text, ("para",))

    def _on_submit(self):
        t = LANGS[self.lang_code]
        text = self.input_text.get("1.0", "end-1c").strip()

        if not self.runner.loaded:
            messagebox.showwarning(t["title"], t["error_model"])
            return
        if not text:
            messagebox.showwarning(t["title"], t["error_empty"])
            return

        beams = self.current_config.num_beams
        self.submit_btn.config_state("disabled")
        self.status_var.set(t["simplifying"])
        self._set_output("")
        self._reset_markdown_toggle()

        prefix = build_prefix(self.current_config)
        self.log.add(f"Запуск упрощения: {len(text)} симв., beams={beams}, prefix='{prefix.strip()}'")
        started = time.monotonic()

        def worker():
            try:
                result = self.runner.simplify(text, beams, prefix=prefix)
            except Exception as e:  # noqa: BLE001
                result = None
                err = str(e)
            else:
                err = None

            def finish():
                elapsed = time.monotonic() - started
                self.submit_btn.config_state("normal")
                if err:
                    self.status_var.set(f"Error: {err}")
                    self.log.add(f"Ошибка упрощения ({elapsed:.1f}s): {err}", level="ERROR")
                    messagebox.showerror(t["title"], err)
                else:
                    self.status_var.set(t["ready"])
                    self._set_output(result)
                    self.log.add(f"Упрощение завершено за {elapsed:.1f}s, получено {len(result)} симв.")

            self.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()


if __name__ == "__main__":
    app = SimplifierApp()
    app.mainloop()
