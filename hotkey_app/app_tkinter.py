"""
Uzbek Text Simplifier — Tkinter desktop app.

Модель грузится напрямую с Hugging Face Hub (репозиторий MODEL_ID),
поэтому веса заливать вручную не нужно — при первом запуске
они скачаются автоматически и закешируются.

Запуск локально:
    pip install -r requirements.txt
    python app_tkinter.py
"""

import logging
import os
import threading
import time
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    _DND_AVAILABLE = True
    _BaseTk = TkinterDnD.Tk
except ImportError:
    _DND_AVAILABLE = False
    _BaseTk = tk.Tk

try:
    import docx as _docx_lib  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from config import SimplifierConfig
from settings_window import SettingsWindow
import global_hotkey
import popup as popup_module
import tray as tray_module

logger = logging.getLogger(__name__)
_LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "hotkey_debug.log")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.StreamHandler(),  # видно в консоли
        logging.FileHandler(_LOG_FILE, encoding="utf-8"),  # остаётся на диске
    ],
)

MODEL_ID = "ismailoviskandar02/uzbek-text-simplifier"  # твой репозиторий модели на HF
PREFIX = "simplify: "
MAX_INPUT_LEN = 256
MAX_NEW_TOKENS = 256

# ---------------------------------------------------------------------------
# Локализация (RU / UZ / EN)
# ---------------------------------------------------------------------------

LANGS = {
    "ru": {
        "flag": "🇷🇺",
        "name": "Русский",
        "title": "🇺🇿 Oʻzbek matnini soddalashtirish",
        "subtitle": "Упрощение узбекского текста с помощью нейросети",
        "description": (
            "Переводит сложный текст (юридический, официальный, государственный "
            "документ) в понятный, простой узбекский язык."
        ),
        "input_label": "Исходный текст",
        "input_placeholder": "Вставьте сложный узбекский текст сюда...",
        "beams_label": "Точность упрощения",
        "beams_fast": "Быстро",
        "beams_medium": "Средне",
        "beams_accurate": "Точно",
        "submit": "✨ Упростить",
        "clear": "🗑️ Очистить",
        "open_file": "📂 Открыть файл",
        "download_docx": "⬇️ Скачать .docx",
        "dnd_hint": "Перетащите .docx или .txt файл сюда, или",
        "output_label": "Упрощённый текст",
        "examples_label": "Примеры",
        "footer": "Сделано с ❤️ для узбекского языка",
        "loading": "Загрузка модели...",
        "ready": "Модель загружена.",
        "simplifying": "Упрощение...",
        "error_empty": "Введите текст для упрощения.",
        "error_model": "Модель ещё не загружена, подождите.",
        "error_file_read": "Не удалось прочитать файл",
        "error_file_type": "Поддерживаются только файлы .docx и .txt",
        "error_no_output": "Сначала получите упрощённый текст.",
        "error_docx_lib": "Для сохранения в .docx нужна библиотека python-docx (pip install python-docx).",
        "saved_ok": "Файл сохранён",
    },
    "uz": {
        "flag": "🇺🇿",
        "name": "Oʻzbek tili",
        "title": "🇺🇿 Oʻzbek matnini soddalashtirish",
        "subtitle": "Sun'iy intellekt yordamida matnni soddalashtirish",
        "description": (
            "Murakkab matnni (huquqiy, rasmiy, davlat hujjati) tushunarli, "
            "oddiy oʻzbek tiliga oʻgiradi."
        ),
        "input_label": "Asl matn",
        "input_placeholder": "Murakkab oʻzbek matnini shu yerga joylashtiring...",
        "beams_label": "Soddalashtirish aniqligi",
        "beams_fast": "Tez",
        "beams_medium": "O'rtacha",
        "beams_accurate": "Aniq",
        "submit": "✨ Soddalashtirish",
        "clear": "🗑️ Tozalash",
        "open_file": "📂 Faylni ochish",
        "download_docx": "⬇️ .docx yuklab olish",
        "dnd_hint": "Bu yerga .docx yoki .txt faylni tashlang, yoki",
        "output_label": "Soddalashtirilgan matn",
        "examples_label": "Namunalar",
        "footer": "Oʻzbek tili uchun ❤️ bilan yaratilgan",
        "loading": "Model yuklanmoqda...",
        "ready": "Model yuklandi.",
        "simplifying": "Soddalashtirilmoqda...",
        "error_empty": "Soddalashtirish uchun matn kiriting.",
        "error_model": "Model hali yuklanmagan, kuting.",
        "error_file_read": "Faylni o'qib bo'lmadi",
        "error_file_type": "Faqat .docx va .txt fayllar qo'llab-quvvatlanadi",
        "error_no_output": "Avval soddalashtirilgan matnni oling.",
        "error_docx_lib": ".docx saqlash uchun python-docx kutubxonasi kerak (pip install python-docx).",
        "saved_ok": "Fayl saqlandi",
    },
    "en": {
        "flag": "🇬🇧",
        "name": "English",
        "title": "🇺🇿 Uzbek Text Simplifier",
        "subtitle": "Simplify Uzbek text with AI",
        "description": (
            "Turns complex text (legal, official, government documents) "
            "into clear, simple Uzbek."
        ),
        "input_label": "Original text",
        "input_placeholder": "Paste complex Uzbek text here...",
        "beams_label": "Simplification accuracy",
        "beams_fast": "Fast",
        "beams_medium": "Medium",
        "beams_accurate": "Accurate",
        "submit": "✨ Simplify",
        "clear": "🗑️ Clear",
        "open_file": "📂 Open file",
        "download_docx": "⬇️ Download .docx",
        "dnd_hint": "Drop a .docx or .txt file here, or",
        "output_label": "Simplified text",
        "examples_label": "Examples",
        "footer": "Made with ❤️ for the Uzbek language",
        "loading": "Loading model...",
        "ready": "Model loaded.",
        "simplifying": "Simplifying...",
        "error_empty": "Please enter text to simplify.",
        "error_model": "Model is still loading, please wait.",
        "error_file_read": "Could not read the file",
        "error_file_type": "Only .docx and .txt files are supported",
        "error_no_output": "Simplify some text first.",
        "error_docx_lib": "Saving .docx requires the python-docx package (pip install python-docx).",
        "saved_ok": "File saved",
    },
}

EXAMPLES = [
    "Oʻzbekiston Respublikasi Vazirlar Mahkamasining qarori bilan tasdiqlangan "
    "Nizomga muvofiq, davlat organlari oʻz vakolatlari doirasida tegishli "
    "chora-tadbirlarni amalga oshiradilar.",
    "Ushbu shartnoma tomonlar oʻrtasida oʻzaro kelishilgan holda tuzilgan boʻlib, "
    "unga muvofiq har bir tomon oʻz majburiyatlarini belgilangan muddatlarda "
    "bajarishi shart.",
]

# ---------------------------------------------------------------------------
# Цветовая палитра (тёмная сине-чёрная тема)
# ---------------------------------------------------------------------------

COLORS = {
    "bg": "#05070d",
    "panel": "#0b0f1a",
    "panel2": "#0f1524",
    "border": "#1e2740",
    "label_bg": "#141b2e",
    "label_fg": "#8fb3ff",
    "title_fg": "#e6edff",
    "text_fg": "#dbe4ff",
    "subdued_fg": "#7d8bb3",
    "input_bg": "#0d1220",
    "input_border": "#243257",
    "accent": "#3b6bff",
    "accent_hover": "#5b8bff",
    "btn_secondary_bg": "#141b2e",
    "btn_secondary_fg": "#c3cdf0",
}


class ModelRunner:
    """Loads and runs the Uzbek text simplification model in the background."""

    def __init__(self, on_loaded=None, on_error=None):
        self.tokenizer = None
        self.model = None
        self.device = "cpu"
        self.loaded = False
        self.on_loaded = on_loaded
        self.on_error = on_error

    def load_async(self):
        threading.Thread(target=self._load, daemon=True).start()

    def _load(self):
        try:
            import torch
            from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

            self.device = "cuda" if torch.cuda.is_available() else "cpu"
            self.tokenizer = AutoTokenizer.from_pretrained(MODEL_ID, subfolder="model2")
            self.model = AutoModelForSeq2SeqLM.from_pretrained(
                MODEL_ID, subfolder="model2"
            ).to(self.device)
            self.model.eval()
            self.loaded = True
            if self.on_loaded:
                self.on_loaded()
        except Exception as e:  # noqa: BLE001
            if self.on_error:
                self.on_error(str(e))

    def simplify(self, text: str, num_beams: int = 4) -> str:
        import torch

        text = (text or "").strip()
        if not text:
            return ""
        inputs = self.tokenizer(
            PREFIX + text,
            return_tensors="pt",
            truncation=True,
            max_length=MAX_INPUT_LEN,
        ).to(self.device)

        with torch.no_grad():
            out = self.model.generate(
                **inputs,
                max_new_tokens=MAX_NEW_TOKENS,
                num_beams=num_beams,
            )

        return self.tokenizer.decode(out[0], skip_special_tokens=True)


class SimplifierApp(_BaseTk):
    def __init__(self):
        super().__init__()
        self.lang_code = "ru"
        self.title(LANGS[self.lang_code]["title"])
        self.geometry("1000x720")
        self.configure(bg=COLORS["bg"])
        self.minsize(760, 560)

        self.runner = ModelRunner(on_loaded=self._model_loaded, on_error=self._model_error)
        self.simplifier_config = SimplifierConfig.load()

        self._build_style()
        self._build_ui()
        self._apply_language()

        self.status_var.set(LANGS[self.lang_code]["loading"])
        self.runner.load_async()

        self._hotkey_handle = None
        self._tray_icon = None
        self._active_popup = None
        self._quitting = False

        self._setup_global_hotkey()
        self._setup_tray()

        # Сворачивание в трей вместо закрытия — иначе процесс убивается
        # вместе с глобальным хоткеем.
        self.protocol("WM_DELETE_WINDOW", self._on_window_close)

    # ------------------------------------------------------------------
    # Global hotkey / clipboard popup / tray
    # ------------------------------------------------------------------
    def _setup_global_hotkey(self) -> None:
        combo = self.simplifier_config.hotkey_combo
        try:
            self._hotkey_handle = global_hotkey.register_hotkey(
                combo, self._on_hotkey_pressed
            )
            logger.info(
                "Хоткей '%s' успешно зарегистрирован (backend=%s). "
                "Приложение готово реагировать на нажатие.",
                combo, global_hotkey.backend_name(),
            )
        except global_hotkey.HotkeyRegistrationError as e:
            logger.error("Hotkey registration failed: %s", e)
            messagebox.showerror(
                "Uzbek Text Simplifier",
                f"Не удалось зарегистрировать глобальный хоткей '{combo}':\n\n{e}\n\n"
                "Приложение продолжит работать без глобального хоткея — "
                "упрощение текста по-прежнему доступно через основное окно.",
            )

    def _setup_tray(self) -> None:
        if not tray_module.is_tray_available():
            logger.warning("pystray/Pillow не установлены — системный трей недоступен.")
            return
        try:
            self._tray_icon = tray_module.TrayIcon(
                root=self,
                on_open=self._restore_from_tray,
                on_settings=self._open_settings,
                on_quit=self._quit_app,
            )
        except Exception as e:  # noqa: BLE001
            logger.error("Не удалось инициализировать трей: %s", e)
            self._tray_icon = None

    def _on_window_close(self) -> None:
        if self._tray_icon is not None:
            tray_module.minimize_to_tray(self, self._tray_icon)
        else:
            # Нет трея — обычное поведение: закрыть приложение полностью,
            # аккуратно освободив хоткей.
            self._quit_app()

    def _restore_from_tray(self) -> None:
        tray_module.restore_from_tray(self)

    def _open_settings(self) -> None:
        self._restore_from_tray()
        self._show_settings_window()

    def _quit_app(self) -> None:
        if self._quitting:
            return
        self._quitting = True
        if self._hotkey_handle is not None:
            self._hotkey_handle.unregister()
        if self._tray_icon is not None:
            self._tray_icon.stop()
        self.destroy()

    # ------------------------------------------------------------------
    # Hotkey callback: выполняется в ФОНОВОМ потоке backend'а хоткея
    # (не в потоке Tkinter!). Здесь мы только захватываем буфер обмена и
    # передаём управление обратно в главный поток через after().
    # ------------------------------------------------------------------
    def _on_hotkey_pressed(self) -> None:
        # ВАЖНО: библиотека 'keyboard' молча проглатывает исключения,
        # возникшие внутри callback'а хоткея (не печатает и не пробрасывает
        # их) — поэтому здесь обязательно ловим и логируем любую ошибку
        # сами, иначе при сбое "ничего не произойдёт" без единого следа.
        try:
            self._on_hotkey_pressed_impl()
        except Exception:  # noqa: BLE001
            logger.exception(
                "Необработанная ошибка в обработчике хоткея — "
                "callback был вызван, но упал внутри."
            )

    def _on_hotkey_pressed_impl(self) -> None:
        logger.info("Хоткей нажат — начинаю захват выделенного текста.")
        t_hotkey = time.monotonic()

        if not self.runner.loaded:
            logger.info("Hotkey pressed, but model is not loaded yet — ignoring.")
            self.after(0, lambda: self._toast("Модель ещё загружается, подождите."))
            return

        try:
            text = global_hotkey.capture_selected_text(timeout=1.0)
        except Exception as e:  # noqa: BLE001
            logger.exception("Ошибка при захвате выделенного текста")
            self.after(0, lambda: self._toast(f"Ошибка захвата буфера: {e}"))
            return

        t_clipboard = time.monotonic()
        clipboard_ms = (t_clipboard - t_hotkey) * 1000

        if not text:
            logger.info(
                "Clipboard did not update within timeout (%.1f ms) — no selection.",
                clipboard_ms,
            )
            self.after(0, lambda: self._toast("Нет выделенного текста."))
            return

        logger.info("Clipboard captured in %.1f ms (%d chars)", clipboard_ms, len(text))

        def show():
            self._show_hotkey_popup(text, t_hotkey, t_clipboard)

        self.after(0, show)

    def _toast(self, message: str) -> None:
        """Короткое уведомление, не открывающее popup (используется для
        случаев вроде 'нет выделенного текста')."""
        if self._tray_icon is not None:
            self._tray_icon.notify(message)
        else:
            self.status_var.set(message)

    def _show_hotkey_popup(self, text: str, t_hotkey: float, t_clipboard: float) -> None:
        beams = self.beams_var.get() if hasattr(self, "beams_var") else 4

        def infer_fn(src_text: str) -> str:
            # Переиспользуем УЖЕ загруженный ModelRunner — новая модель не создаётся.
            t_infer_start = time.monotonic()
            result = self.runner.simplify(src_text, beams)
            t_infer_end = time.monotonic()

            total_ms = (t_infer_end - t_hotkey) * 1000
            clipboard_ms = (t_clipboard - t_hotkey) * 1000
            infer_ms = (t_infer_end - t_infer_start) * 1000
            overhead_ms = total_ms - infer_ms

            logger.info(
                "[hotkey timing] total=%.1fms clipboard_capture=%.1fms "
                "inference=%.1fms other_overhead=%.1fms chars_in=%d",
                total_ms, clipboard_ms, infer_ms, overhead_ms, len(src_text),
            )
            return result

        self._active_popup = popup_module.show_popup(
            self, text, infer_fn, on_close=lambda: setattr(self, "_active_popup", None)
        )

    def _show_settings_window(self) -> None:
        def on_saved(new_config: SimplifierConfig) -> None:
            old_combo = self.simplifier_config.hotkey_combo
            self.simplifier_config = new_config
            if new_config.hotkey_combo != old_combo:
                self._reregister_hotkey(new_config.hotkey_combo)

        SettingsWindow(
            self,
            lang_code=self.lang_code,
            on_saved=on_saved,
            set_status=self.status_var.set,
        )

    def _reregister_hotkey(self, new_combo: str) -> None:
        if self._hotkey_handle is not None:
            self._hotkey_handle.unregister()
            self._hotkey_handle = None
        try:
            self._hotkey_handle = global_hotkey.register_hotkey(
                new_combo, self._on_hotkey_pressed
            )
            self.status_var.set(f"Хоткей обновлён: {new_combo}")
        except global_hotkey.HotkeyRegistrationError as e:
            logger.error("Hotkey re-registration failed: %s", e)
            messagebox.showerror(
                "Uzbek Text Simplifier",
                f"Не удалось зарегистрировать новый хоткей '{new_combo}':\n\n{e}",
            )

    # ------------------------------------------------------------------
    # Styling
    # ------------------------------------------------------------------
    def _build_style(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        style.configure(
            "TFrame", background=COLORS["bg"]
        )
        style.configure(
            "Panel.TFrame", background=COLORS["panel"]
        )
        style.configure(
            "TLabel", background=COLORS["bg"], foreground=COLORS["text_fg"],
            font=("Segoe UI", 10),
        )
        style.configure(
            "Title.TLabel", background=COLORS["bg"], foreground=COLORS["title_fg"],
            font=("Segoe UI", 20, "bold"),
        )
        style.configure(
            "Subtitle.TLabel", background=COLORS["bg"], foreground=COLORS["label_fg"],
            font=("Segoe UI", 12, "bold"),
        )
        style.configure(
            "Desc.TLabel", background=COLORS["bg"], foreground=COLORS["subdued_fg"],
            font=("Segoe UI", 10), wraplength=900, justify="center",
        )
        style.configure(
            "FieldLabel.TLabel", background=COLORS["bg"], foreground=COLORS["label_fg"],
            font=("Segoe UI", 10, "bold"),
        )
        style.configure(
            "Status.TLabel", background=COLORS["bg"], foreground=COLORS["subdued_fg"],
            font=("Segoe UI", 9, "italic"),
        )
        style.configure(
            "Footer.TLabel", background=COLORS["bg"], foreground=COLORS["subdued_fg"],
            font=("Segoe UI", 9),
        )

        style.configure(
            "Accent.TButton",
            background=COLORS["accent"],
            foreground="#ffffff",
            font=("Segoe UI", 10, "bold"),
            padding=8,
            borderwidth=0,
        )
        style.map(
            "Accent.TButton",
            background=[("active", COLORS["accent_hover"])],
        )

        style.configure(
            "Secondary.TButton",
            background=COLORS["btn_secondary_bg"],
            foreground=COLORS["btn_secondary_fg"],
            font=("Segoe UI", 10),
            padding=8,
            borderwidth=0,
        )
        style.map(
            "Secondary.TButton",
            background=[("active", COLORS["panel2"])],
        )

        style.configure(
            "TScale", background=COLORS["bg"], troughcolor=COLORS["input_bg"],
        )

        style.configure(
            "TCombobox",
            fieldbackground=COLORS["input_bg"],
            background=COLORS["input_bg"],
            foreground=COLORS["text_fg"],
            arrowcolor=COLORS["text_fg"],
        )

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------
    def _build_ui(self):
        root = ttk.Frame(self, style="TFrame", padding=16)
        root.pack(fill="both", expand=True)

        # Language picker row
        lang_row = ttk.Frame(root, style="TFrame")
        lang_row.pack(fill="x")

        self.lang_var = tk.StringVar(value=self.lang_code)
        lang_display = [f"{v['flag']} {v['name']}" for v in LANGS.values()]
        self._lang_keys = list(LANGS.keys())

        self.lang_combo = ttk.Combobox(
            lang_row, values=lang_display, state="readonly", width=16
        )
        self.lang_combo.current(self._lang_keys.index(self.lang_code))
        self.lang_combo.pack(side="right")
        self.lang_combo.bind("<<ComboboxSelected>>", self._on_language_change)

        # Header
        self.title_label = ttk.Label(root, text="", style="Title.TLabel", anchor="center")
        self.title_label.pack(pady=(12, 2), fill="x")

        self.subtitle_label = ttk.Label(root, text="", style="Subtitle.TLabel", anchor="center")
        self.subtitle_label.pack(fill="x")

        self.desc_label = ttk.Label(root, text="", style="Desc.TLabel", anchor="center")
        self.desc_label.pack(pady=(4, 12), fill="x")

        # Main content: two columns
        content = ttk.Frame(root, style="TFrame")
        content.pack(fill="both", expand=True)
        content.columnconfigure(0, weight=1)
        content.columnconfigure(1, weight=1)
        content.rowconfigure(0, weight=1)

        # Left column: input
        left = ttk.Frame(content, style="TFrame")
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 8))

        self.input_label = ttk.Label(left, text="", style="FieldLabel.TLabel")
        self.input_label.pack(anchor="w")

        self.input_text = tk.Text(
            left, height=12, wrap="word",
            bg=COLORS["input_bg"], fg=COLORS["text_fg"],
            insertbackground=COLORS["text_fg"],
            relief="flat", highlightthickness=1,
            highlightbackground=COLORS["input_border"],
            highlightcolor=COLORS["accent"],
            font=("Segoe UI", 11), padx=8, pady=8,
        )
        self.input_text.pack(fill="both", expand=True, pady=(4, 4))

        if _DND_AVAILABLE:
            self.input_text.drop_target_register(DND_FILES)
            self.input_text.dnd_bind("<<Drop>>", self._on_file_drop)

        file_row = ttk.Frame(left, style="TFrame")
        file_row.pack(fill="x", pady=(0, 8))

        self.dnd_hint_label = ttk.Label(file_row, text="", style="Status.TLabel")
        self.dnd_hint_label.pack(side="left")
        if not _DND_AVAILABLE:
            self.dnd_hint_label.pack_forget()

        self.open_file_btn = ttk.Button(
            file_row, text="", style="Secondary.TButton", command=self._on_open_file
        )
        self.open_file_btn.pack(side="left", padx=(6, 0))

        self.beams_label = ttk.Label(left, text="", style="FieldLabel.TLabel")
        self.beams_label.pack(anchor="w")

        self.beams_var = tk.IntVar(value=4)
        beams_scale = ttk.Scale(
            left, from_=1, to=8, orient="horizontal",
            variable=self.beams_var, command=self._on_beams_change,
        )
        beams_scale.pack(fill="x", pady=(4, 4))

        self.beams_value_label = ttk.Label(left, text="", style="Status.TLabel")
        self.beams_value_label.pack(anchor="e")

        btn_row = ttk.Frame(left, style="TFrame")
        btn_row.pack(fill="x", pady=(8, 0))

        self.clear_btn = ttk.Button(
            btn_row, text="", style="Secondary.TButton", command=self._on_clear
        )
        self.clear_btn.pack(side="left", padx=(0, 8))

        self.submit_btn = ttk.Button(
            btn_row, text="", style="Accent.TButton", command=self._on_submit
        )
        self.submit_btn.pack(side="left")

        # Examples
        self.examples_label = ttk.Label(left, text="", style="FieldLabel.TLabel")
        self.examples_label.pack(anchor="w", pady=(12, 4))

        self.examples_frame = ttk.Frame(left, style="TFrame")
        self.examples_frame.pack(fill="x")
        self.example_buttons = []
        for ex in EXAMPLES:
            btn = tk.Button(
                self.examples_frame,
                text=(ex[:60] + "...") if len(ex) > 60 else ex,
                bg=COLORS["panel2"], fg=COLORS["subdued_fg"],
                activebackground=COLORS["panel"], activeforeground=COLORS["text_fg"],
                relief="flat", anchor="w", justify="left", wraplength=420,
                font=("Segoe UI", 9), padx=8, pady=6,
                command=lambda t=ex: self._use_example(t),
            )
            btn.pack(fill="x", pady=2)
            self.example_buttons.append(btn)

        # Right column: output
        right = ttk.Frame(content, style="TFrame")
        right.grid(row=0, column=1, sticky="nsew", padx=(8, 0))

        self.output_label = ttk.Label(right, text="", style="FieldLabel.TLabel")
        self.output_label.pack(anchor="w")

        self.output_text = tk.Text(
            right, height=12, wrap="word",
            bg=COLORS["input_bg"], fg=COLORS["text_fg"],
            insertbackground=COLORS["text_fg"],
            relief="flat", highlightthickness=1,
            highlightbackground=COLORS["input_border"],
            font=("Segoe UI", 11), padx=8, pady=8, state="disabled",
        )
        self.output_text.pack(fill="both", expand=True, pady=(4, 8))

        self.download_btn = ttk.Button(
            right, text="", style="Secondary.TButton", command=self._on_download_docx
        )
        self.download_btn.pack(anchor="e")

        # Status + footer
        self.status_var = tk.StringVar(value="")
        self.status_label = ttk.Label(root, textvariable=self.status_var, style="Status.TLabel")
        self.status_label.pack(pady=(8, 0))

        self.footer_label = ttk.Label(root, text="", style="Footer.TLabel", anchor="center")
        self.footer_label.pack(pady=(8, 0), fill="x")

    # ------------------------------------------------------------------
    # Language handling
    # ------------------------------------------------------------------
    def _on_language_change(self, event=None):
        idx = self.lang_combo.current()
        self.lang_code = self._lang_keys[idx]
        self._apply_language()

    def _apply_language(self):
        t = LANGS[self.lang_code]
        self.title(t["title"])
        self.title_label.config(text=t["title"])
        self.subtitle_label.config(text=t["subtitle"])
        self.desc_label.config(text=t["description"])
        self.input_label.config(text=t["input_label"])
        self.beams_label.config(text=t["beams_label"])
        self.beams_value_label.config(
            text=f"{self._beams_hint(self.beams_var.get())} ({self.beams_var.get()})"
        )
        self.submit_btn.config(text=t["submit"])
        self.clear_btn.config(text=t["clear"])
        self.open_file_btn.config(text=t["open_file"])
        self.dnd_hint_label.config(text=t["dnd_hint"])
        self.download_btn.config(text=t["download_docx"])
        self.output_label.config(text=t["output_label"])
        self.examples_label.config(text=t["examples_label"])
        self.footer_label.config(text=t["footer"])

        placeholder = t["input_placeholder"]
        self._set_placeholder(placeholder)

        if self.runner.loaded:
            self.status_var.set(t["ready"])
        elif self.status_var.get():
            self.status_var.set(t["loading"])

    def _set_placeholder(self, placeholder):
        # simple placeholder: only set if the field is currently empty
        current = self.input_text.get("1.0", "end-1c").strip()
        if not current:
            pass  # tk.Text has no native placeholder; label above the field carries the hint
        self._placeholder_text = placeholder

    # ------------------------------------------------------------------
    # Model callbacks
    # ------------------------------------------------------------------
    def _model_loaded(self):
        self.after(0, lambda: self.status_var.set(LANGS[self.lang_code]["ready"]))

    def _model_error(self, err):
        def show():
            self.status_var.set(f"Error: {err}")
            messagebox.showerror("Model load error", err)
        self.after(0, show)

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------
    def _beams_hint(self, v: int) -> str:
        t = LANGS[self.lang_code]
        if v <= 2:
            return t["beams_fast"]
        if v <= 5:
            return t["beams_medium"]
        return t["beams_accurate"]

    def _on_beams_change(self, value):
        v = int(round(float(value)))
        self.beams_var.set(v)
        self.beams_value_label.config(text=f"{self._beams_hint(v)} ({v})")

    def _use_example(self, text):
        self.input_text.delete("1.0", "end")
        self.input_text.insert("1.0", text)

    def _on_clear(self):
        self.input_text.delete("1.0", "end")
        self._set_output("")

    # ------------------------------------------------------------------
    # File import (drag & drop + open dialog)
    # ------------------------------------------------------------------
    def _read_text_file(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        if ext == ".docx":
            doc = _docx_lib.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError("unsupported_type")

    def _load_file_into_input(self, path: str):
        t = LANGS[self.lang_code]
        ext = os.path.splitext(path)[1].lower()
        if ext not in (".txt", ".docx"):
            messagebox.showwarning(t["title"], t["error_file_type"])
            return
        if ext == ".docx" and not _DOCX_AVAILABLE:
            messagebox.showerror(t["title"], t["error_docx_lib"])
            return
        try:
            text = self._read_text_file(path)
        except Exception as e:  # noqa: BLE001
            messagebox.showerror(t["title"], f"{t['error_file_read']}: {e}")
            return
        self.input_text.delete("1.0", "end")
        self.input_text.insert("1.0", text)

    def _on_open_file(self):
        path = filedialog.askopenfilename(
            filetypes=[("Text / Word", "*.txt *.docx"), ("All files", "*.*")]
        )
        if path:
            self._load_file_into_input(path)

    def _on_file_drop(self, event):
        # tkinterdnd2 gives a Tcl list; braces wrap paths containing spaces
        paths = self.tk.splitlist(event.data)
        if paths:
            self._load_file_into_input(paths[0])

    # ------------------------------------------------------------------
    # File export (.docx)
    # ------------------------------------------------------------------
    def _on_download_docx(self):
        t = LANGS[self.lang_code]
        text = self.output_text.get("1.0", "end-1c").strip()
        if not text:
            messagebox.showwarning(t["title"], t["error_no_output"])
            return
        if not _DOCX_AVAILABLE:
            messagebox.showerror(t["title"], t["error_docx_lib"])
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            initialfile="simplified.docx",
        )
        if not path:
            return
        try:
            doc = _docx_lib.Document()
            for line in text.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
        except Exception as e:  # noqa: BLE001
            messagebox.showerror(t["title"], f"{t['error_file_read']}: {e}")
            return
        self.status_var.set(f"{t['saved_ok']}: {os.path.basename(path)}")

    def _set_output(self, text):
        self.output_text.config(state="normal")
        self.output_text.delete("1.0", "end")
        self.output_text.insert("1.0", text)
        self.output_text.config(state="disabled")

    def _on_submit(self):
        t = LANGS[self.lang_code]
        text = self.input_text.get("1.0", "end-1c").strip()

        if not self.runner.loaded:
            messagebox.showwarning(t["title"], t["error_model"])
            return
        if not text:
            messagebox.showwarning(t["title"], t["error_empty"])
            return

        beams = self.beams_var.get()
        self.submit_btn.config(state="disabled")
        self.status_var.set(t["simplifying"])
        self._set_output("")

        def worker():
            try:
                result = self.runner.simplify(text, beams)
            except Exception as e:  # noqa: BLE001
                result = None
                err = str(e)
            else:
                err = None

            def finish():
                self.submit_btn.config(state="normal")
                if err:
                    self.status_var.set(f"Error: {err}")
                    messagebox.showerror(t["title"], err)
                else:
                    self.status_var.set(t["ready"])
                    self._set_output(result)

            self.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()


if __name__ == "__main__":
    app = SimplifierApp()
    app.mainloop()
